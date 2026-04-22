"""Apply Leeds template heading hierarchy and styling to the rendered docx.

Pandoc renders our markdown with one-too-deep heading levels relative to
the Leeds COMP3931 template convention:
  template: H1 = chapter, H2 = section, H3 = subsection, H4 = sub-subsection
  pandoc:   H1 = report title, H2 = chapter, H3 = section, H4 = subsection

This script:
  1. Enforces Leeds spec: 11pt body, 1.5 line spacing, 2.5cm margins.
  2. Tightens heading spacing to fit the 30-page body limit.
  3. Promotes every chapter heading to Heading 1.
  4. Promotes every X.Y section heading from Heading 3 to Heading 2.
  5. Promotes every X.Y.Z subsection heading from Heading 4 to Heading 3.
  6. Centres the title-page block.
  7. Adds page breaks before each chapter (Heading 1).
  8. Inserts all 10 figure images before their italic captions
     (pandoc does not embed raw <img> HTML tags).
  9. Italicises figure captions and centres them.
 10. Centres tables and applies clean grid styling.
"""
import re
import sys
from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm, Emu, Inches, RGBColor
from docx.enum.text import WD_BREAK, WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from copy import deepcopy

REPORT = Path(__file__).resolve().parent.parent / "docs" / "report"
# Pandoc renders the intermediate docx; this script post-processes it in place.
SRC = REPORT / "Final_Report_Draft_template.docx"
DST = REPORT / "Final_Report_Draft_template.docx"
# Final PDF is rendered from the docx into the canonical primary path
# Final_Report_Draft.pdf (no separate _template.pdf duplicate).


# Patterns that should be Heading 1 (chapter-level)
H1_PATTERNS = [
    r"^Chapter \d+:",
    r"^Appendix [AB]:",
    r"^List of References$",
    r"^Summary$",
    r"^Acknowledgements$",
    r"^Declaration$",
    r"^Deliverables$",
    r"^Table of Contents$",
    r"^List of Figures$",
    r"^List of Tables$",
]

# Patterns for section level (X.Y - currently Heading 3, should be Heading 2)
H2_PATTERN = re.compile(r"^(?:\d+\.\d+|[AB]\.\d+|B\.\d+)\s+\S")

# Patterns for subsection level (X.Y.Z - currently Heading 4, should be Heading 3)
H3_PATTERN = re.compile(r"^(?:\d+\.\d+\.\d+|[AB]\.\d+\.\d+)\s+\S")

# Patterns for the main report title (should become Title style)
TITLE_TEXT = "Audit-Ready Policy Copilot"


def matches_h1(text: str) -> bool:
    return any(re.match(p, text) for p in H1_PATTERNS)


def add_page_break_before(paragraph):
    """Insert a page break before the given paragraph by adding pageBreakBefore property."""
    pPr = paragraph._element.get_or_add_pPr()
    pageBreakBefore = OxmlElement("w:pageBreakBefore")
    pPr.append(pageBreakBefore)


def split_sections_for_page_numbering(doc, chapter_one_paragraph):
    """
    Configure two sections for proper Leeds page numbering:
    - Section 1 (preliminaries, before Ch 1): lower-roman numerals
    - Section 2 (body + refs + appendices, from Ch 1): decimal (Arabic) restarting at 1

    This works by:
    1. Inserting a section break (sectPr) inside the LAST paragraph BEFORE
       Chapter 1, configured for Roman numerals (this defines Section 1).
    2. Setting the document-level body sectPr to Arabic restart-at-1
       (this defines Section 2 = everything after the section break).
    """
    # Find the paragraph immediately before Chapter 1 by XML element identity
    body_xml = doc.element.body
    paras = list(doc.paragraphs)
    ch1_idx = None
    target_elem = chapter_one_paragraph._element
    for i, p in enumerate(paras):
        if p._element is target_elem:
            ch1_idx = i
            break
    if ch1_idx is None or ch1_idx == 0:
        return False
    last_prelim = paras[ch1_idx - 1]

    # Build sectPr for section 1 (preliminaries: lower roman)
    sect1 = OxmlElement("w:sectPr")
    pgSz1 = OxmlElement("w:pgSz")
    pgSz1.set(qn("w:w"), "11906")
    pgSz1.set(qn("w:h"), "16838")
    sect1.append(pgSz1)
    pgMar1 = OxmlElement("w:pgMar")
    for k, v in [("top", "1417"), ("right", "1417"), ("bottom", "1417"),
                 ("left", "1417"), ("header", "708"), ("footer", "708"), ("gutter", "0")]:
        pgMar1.set(qn(f"w:{k}"), v)
    sect1.append(pgMar1)
    pgNumType1 = OxmlElement("w:pgNumType")
    pgNumType1.set(qn("w:fmt"), "lowerRoman")
    pgNumType1.set(qn("w:start"), "1")
    sect1.append(pgNumType1)
    sect1Type = OxmlElement("w:type")
    sect1Type.set(qn("w:val"), "nextPage")
    sect1.append(sect1Type)

    # Attach sect1 sectPr to the last prelim paragraph
    pPr = last_prelim._element.get_or_add_pPr()
    pPr.append(sect1)

    # Now configure document body sectPr for section 2 (Arabic restart)
    body_sectPr = body_xml.find(qn("w:sectPr"))
    if body_sectPr is None:
        body_sectPr = OxmlElement("w:sectPr")
        body_xml.append(body_sectPr)
    # Remove existing pgNumType in body sectPr
    for e in body_sectPr.findall(qn("w:pgNumType")):
        body_sectPr.remove(e)
    # Add Arabic restart-at-1
    pgNumType2 = OxmlElement("w:pgNumType")
    pgNumType2.set(qn("w:fmt"), "decimal")
    pgNumType2.set(qn("w:start"), "1")
    body_sectPr.append(pgNumType2)

    return True


# Figure caption text -> (image filename, width in inches) mapping
# Caption matched by case-insensitive prefix to handle minor edits
FIGURE_MAP = [
    # (caption prefix, image filename, width in inches) - tightened to stay within 30-page body limit
    ("Figure 1.1: PRISMA", "fig_prisma.png", 3.5),
    ("Figure 2.0: Gantt chart", "fig_gantt.png", 4.5),
    ("Figure 2.1: End-to-end pipeline", "fig_data_flow.png", 4.5),
    ("Figure 4.1: Grouped bar chart", "fig_baselines.png", 4.0),
    ("Figure 4.2: Retrieval quality", "fig_retrieval.png", 4.0),
    ("Figure 4.3: Groundedness", "fig_groundedness.png", 4.0),
    ("Figure 4.4: Coverage", "fig_tradeoff.png", 3.5),
    ("Figure B.1: Answerable query result", "screenshot_answerable_query.png", 5.0),
    ("Figure B.2: Unanswerable query showing", "screenshot_unanswerable_query.png", 5.0),
    ("Figure B.3: Contradiction query showing", "screenshot_contradiction_query.png", 5.0),
]


def insert_figures_before_captions(doc, figures_dir):
    """For each italic figure caption, insert the corresponding image
    as a centered paragraph immediately before it. Strip surrounding
    empty paragraphs that pandoc inserted from div-block boundaries
    to keep the figure tightly attached to its caption."""
    inserted = 0
    skipped = 0
    seen = set()

    body = doc.element.body

    for p in list(doc.paragraphs):
        text = p.text.strip()
        if not text or not text.startswith("Figure"):
            continue
        matched = None
        for prefix, fname, width in FIGURE_MAP:
            if text.startswith(prefix) and prefix not in seen:
                matched = (fname, width)
                seen.add(prefix)
                break
        if not matched:
            continue
        fname, width = matched
        img_path = figures_dir / fname
        if not img_path.exists():
            print(f"  WARNING: image not found: {img_path}")
            skipped += 1
            continue

        # Remove empty paragraphs immediately before this caption
        cap_elem = p._element
        prev = cap_elem.getprevious()
        removed = 0
        while prev is not None and prev.tag.endswith("}p"):
            prev_text = "".join(prev.itertext()).strip()
            # Don't remove if it has substantive text or contains an image
            if prev_text or prev.findall(".//" + qn("w:drawing")):
                break
            to_remove = prev
            prev = prev.getprevious()
            body.remove(to_remove)
            removed += 1

        # Insert a new paragraph BEFORE this caption with the image
        new_p = p.insert_paragraph_before("")
        new_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        # Tighten spacing on image paragraph
        new_p.paragraph_format.space_before = Pt(2)
        new_p.paragraph_format.space_after = Pt(0)
        run = new_p.add_run()
        try:
            run.add_picture(str(img_path), width=Inches(width))
            inserted += 1
        except Exception as e:
            print(f"  WARNING: failed to insert {fname}: {e}")
            skipped += 1

        # Style the caption: italic + centered + smaller, tight spacing
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(4)
        for run in p.runs:
            run.font.italic = True
            run.font.size = Pt(10)

    print(f"Inserted {inserted} figure images, skipped {skipped}")


def style_tables(doc):
    """Apply consistent professional styling to all tables: centred,
    grid borders, header row shaded, compact 9pt cell text."""
    for table in doc.tables:
        # Centre the table on the page
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        # Apply grid style if available
        try:
            table.style = "Table Grid"
        except KeyError:
            pass
        # Style cells
        for row_idx, row in enumerate(table.rows):
            is_header = row_idx == 0
            for cell in row.cells:
                # Shade header row
                if is_header:
                    tcPr = cell._element.get_or_add_tcPr()
                    shd = OxmlElement("w:shd")
                    shd.set(qn("w:val"), "clear")
                    shd.set(qn("w:color"), "auto")
                    shd.set(qn("w:fill"), "2C3E50")
                    tcPr.append(shd)
                # Reduce cell padding via tcMar (cell margins)
                tcPr = cell._element.get_or_add_tcPr()
                tcMar = OxmlElement("w:tcMar")
                for side, val in [("top", "40"), ("bottom", "40"), ("left", "80"), ("right", "80")]:
                    mar = OxmlElement(f"w:{side}")
                    mar.set(qn("w:w"), val)
                    mar.set(qn("w:type"), "dxa")
                    tcMar.append(mar)
                tcPr.append(tcMar)
                for para in cell.paragraphs:
                    para.paragraph_format.line_spacing = 1.0
                    para.paragraph_format.space_before = Pt(0)
                    para.paragraph_format.space_after = Pt(0)
                    for run in para.runs:
                        run.font.size = Pt(9)
                        if is_header:
                            run.font.bold = True
                            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)


def style_table_captions(doc):
    """Bold and slightly larger 'Table N.N:' captions; left-aligned."""
    for p in doc.paragraphs:
        text = p.text.strip()
        if re.match(r"^Table \d+\.\d+:", text) or re.match(r"^Table \d+\.\d+a?:", text):
            for run in p.runs:
                run.font.bold = True
                run.font.size = Pt(10.5)


def enforce_leeds_spec(doc):
    """Enforce Leeds spec: 11pt body, 1.5 line spacing, 2.5cm margins, tight headings."""
    # Margins: 2.5cm
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # Normal body: 11pt, 1.5 line spacing (Leeds spec), minimal paragraph spacing
    normal = doc.styles["Normal"]
    normal.font.size = Pt(11)
    normal.font.name = "Times New Roman"
    pf = normal.paragraph_format
    pf.line_spacing = 1.5
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)  # tightened from 2pt
    pf.widow_control = True

    # Tighten Compact / List Paragraph styles too
    for sname in ["Compact", "List Paragraph"]:
        try:
            cs = doc.styles[sname]
            cs.font.size = Pt(11)
            cspf = cs.paragraph_format
            cspf.line_spacing = 1.5
            cspf.space_before = Pt(0)
            cspf.space_after = Pt(0)
        except (KeyError, AttributeError):
            pass

    # Compact code blocks
    for sname in ["Source Code", "Verbatim Char", "Code"]:
        try:
            cs = doc.styles[sname]
            cs.font.size = Pt(9)
            cs.paragraph_format.line_spacing = 1.1
            cs.paragraph_format.space_before = Pt(1)
            cs.paragraph_format.space_after = Pt(1)
        except (KeyError, AttributeError):
            pass

    # Heading 1: chapter level — 15pt, tight
    # The Leeds template's Heading 1 has built-in pageBreakBefore which
    # forces every chapter to a new page. We strip that here so chapters
    # 2-5 flow continuously, saving body-page budget.
    h1 = doc.styles["Heading 1"]
    h1.font.size = Pt(15)
    h1.font.bold = True
    h1pf = h1.paragraph_format
    h1pf.space_before = Pt(8)
    h1pf.space_after = Pt(4)
    h1pf.line_spacing = 1.15
    h1pf.keep_with_next = True
    # Remove built-in pageBreakBefore from the H1 style itself
    h1_pPr = h1.element.find(qn("w:pPr"))
    if h1_pPr is not None:
        for el in h1_pPr.findall(qn("w:pageBreakBefore")):
            h1_pPr.remove(el)

    # Heading 2: section — 12pt, tight
    h2 = doc.styles["Heading 2"]
    h2.font.size = Pt(12)
    h2.font.bold = True
    h2pf = h2.paragraph_format
    h2pf.space_before = Pt(4)
    h2pf.space_after = Pt(2)
    h2pf.line_spacing = 1.15
    h2pf.keep_with_next = True

    # Heading 3: subsection — 11pt bold
    h3 = doc.styles["Heading 3"]
    h3.font.size = Pt(11)
    h3.font.bold = True
    h3pf = h3.paragraph_format
    h3pf.space_before = Pt(3)
    h3pf.space_after = Pt(1)
    h3pf.line_spacing = 1.15
    h3pf.keep_with_next = True

    # Heading 4: sub-subsection — 11pt italic
    try:
        h4 = doc.styles["Heading 4"]
        h4.font.size = Pt(11)
        h4.font.bold = False
        h4.font.italic = True
        h4pf = h4.paragraph_format
        h4pf.space_before = Pt(2)
        h4pf.space_after = Pt(1)
        h4pf.line_spacing = 1.15
        h4pf.keep_with_next = True
    except KeyError:
        pass


def remove_consecutive_empty_paragraphs(doc):
    """Collapse runs of >1 consecutive empty paragraphs to a single one
    to recover wasted vertical space introduced by pandoc div boundaries."""
    body = doc.element.body
    removed = 0
    paras = list(doc.paragraphs)
    prev_empty = False
    for p in paras:
        text = "".join(p._element.itertext()).strip()
        has_image = bool(p._element.findall(".//" + qn("w:drawing")))
        is_empty = not text and not has_image
        if is_empty and prev_empty:
            try:
                body.remove(p._element)
                removed += 1
            except Exception:
                pass
        prev_empty = is_empty
    print(f"Removed {removed} duplicate empty paragraphs")


def main() -> int:
    doc = Document(str(SRC))
    enforce_leeds_spec(doc)

    # Insert figure images before each italic caption
    figures_dir = REPORT / "figures"
    insert_figures_before_captions(doc, figures_dir)

    # Style tables and table captions
    style_tables(doc)
    style_table_captions(doc)

    # Remove duplicate consecutive empty paragraphs
    remove_consecutive_empty_paragraphs(doc)

    promoted_h1 = 0
    promoted_h2 = 0
    promoted_h3 = 0
    title_demoted = False
    seen_declaration = False
    chapter_one_para = None

    for p in doc.paragraphs:
        text = p.text.strip()
        if not text:
            continue
        current_style = p.style.name if p.style else ""

        if text == "Declaration":
            seen_declaration = True

        # Demote the main report title from Heading 1 to Normal so it doesn't
        # force a page break and sits on the title page with everything else.
        if text == TITLE_TEXT and not title_demoted:
            p.style = doc.styles["Normal"]
            # Make it big and bold via direct run formatting
            for run in p.runs:
                run.font.size = Pt(22)
                run.font.bold = True
            title_demoted = True
            continue

        # Demote the subtitle (currently Heading 2) to Normal too
        if not seen_declaration and "Heading 2" in current_style:
            # If this is the subtitle (Evidence-Grounded...) demote it
            if text.startswith("Evidence-Grounded"):
                p.style = doc.styles["Normal"]
                for run in p.runs:
                    run.font.size = Pt(14)
                    run.font.italic = True
                continue

        # Promote H2 -> H1 for chapter-level headings (only after Declaration)
        if "Heading 2" in current_style and matches_h1(text):
            p.style = doc.styles["Heading 1"]
            # Force a page break for major front-matter, References, and Appendices.
            # Chapters 2-5 flow continuously (we strip H1 style page-break-before
            # in enforce_leeds_spec) to stay within the 30-page body limit.
            wants_page_break = (
                text == "Declaration"
                or text == "Deliverables"
                or text == "Summary"
                or text == "Acknowledgements"
                or text == "Table of Contents"
                or text == "List of Figures"
                or text == "List of Tables"
                or text.startswith("Chapter 1:")
                or text == "List of References"
                or text.startswith("Appendix ")
            )
            if wants_page_break:
                add_page_break_before(p)
            if text.startswith("Chapter 1:") and chapter_one_para is None:
                chapter_one_para = p
            promoted_h1 += 1
            continue

        # Promote H3 -> H2 for X.Y section headings
        if "Heading 3" in current_style and H2_PATTERN.match(text):
            p.style = doc.styles["Heading 2"]
            promoted_h2 += 1
            continue

        # Promote H4 -> H3 for X.Y.Z subsection headings
        if "Heading 4" in current_style and H3_PATTERN.match(text):
            p.style = doc.styles["Heading 3"]
            promoted_h3 += 1
            continue

    # Centre the title-page elements (everything before "Declaration")
    for p in doc.paragraphs:
        if p.text.strip() == "Declaration":
            break
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Insert section break for Roman -> Arabic page numbering
    section_split = False
    if chapter_one_para is not None:
        section_split = split_sections_for_page_numbering(doc, chapter_one_para)

    doc.save(str(DST))
    print(f"Promoted {promoted_h1} headings to Heading 1 (chapter level)")
    print(f"Promoted {promoted_h2} headings to Heading 2 (section level)")
    print(f"Promoted {promoted_h3} headings to Heading 3 (subsection level)")
    print(f"Title demoted to Normal: {title_demoted}")
    print(f"Section split for page numbering: {section_split}")
    print(f"Wrote: {DST}")
    return 0


if __name__ == "__main__":
    sys.exit(main())
